#!/usr/bin/env python
"""
Extract text from uploaded documents and map it to IDX tickers.

Supported formats:
- PDF, DOCX, XLSX/XLS, CSV, TXT, MD, HTML, JSON

The output is a research triage layer. It highlights possible ticker/entity
mentions, sentiment, impact tags, risk flags, and snippets that should be
checked against the original document.
"""

import argparse
import csv
import html
import json
import re
import sys
from collections import Counter
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_WATCHLIST = ROOT / "data" / "watchlist.idx.csv"
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md", ".html", ".htm", ".json"}

sys.path.insert(0, str(Path(__file__).resolve().parent))
try:
    from process_news import IMPACT_TAG_TERMS, RISK_FLAG_TERMS, normalize_text, score_article, tags_for_text
except Exception:
    IMPACT_TAG_TERMS = {}
    RISK_FLAG_TERMS = {}

    def normalize_text(value: Any) -> str:
        return re.sub(r"\s+", " ", "" if value is None else str(value)).strip()

    def score_article(title: str, snippet: str):
        return 0.0, "neutral", [], [], [], [], 0.25

    def tags_for_text(text: str, taxonomy: Dict[str, Sequence[str]]) -> List[str]:
        return []


@dataclass
class DocumentHit:
    ticker: str
    company: str
    source_file: str
    file_type: str
    mention_count: int
    matched_terms: str
    sentiment_score: float
    sentiment_label: str
    impact_tags: str
    risk_flags: str
    confidence: float
    snippets: str
    processed_at: str


def utc_now_text() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")


def read_watchlist(path: Path) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append({key: normalize_text(value) for key, value in row.items()})
    return rows


def split_aliases(value: str) -> List[str]:
    return [part.strip() for part in re.split(r"[;|,]", value or "") if part.strip()]


def ticker_base(ticker: str) -> str:
    return normalize_text(ticker).upper().replace(".JK", "")


def clean_entity_term(value: str) -> str:
    value = normalize_text(value)
    value = re.sub(r"\b(pt|tbk|terbuka|persero)\b\.?", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def profile_terms(profile: Dict[str, str]) -> List[str]:
    ticker = ticker_base(profile.get("ticker", ""))
    company = clean_entity_term(profile.get("company", ""))
    aliases = [clean_entity_term(alias) for alias in split_aliases(profile.get("aliases", ""))]
    terms: List[str] = []
    for term in [ticker, profile.get("ticker", ""), company, *aliases]:
        term = normalize_text(term)
        if not term:
            continue
        if len(term) < 3:
            continue
        if term.lower() in {"pt", "tbk", "persero"}:
            continue
        if term.lower() not in {item.lower() for item in terms}:
            terms.append(term)
    return terms[:10]


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages[:120], start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            pages.append(f"[page {index}] {text}")
    return "\n".join(pages)


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx belum terinstall; jalankan pip install python-docx") from exc

    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_spreadsheet(path: Path) -> str:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    parts = []
    for sheet_name, frame in sheets.items():
        frame = frame.fillna("")
        parts.append(f"[sheet {sheet_name}]")
        for _, row in frame.head(2500).iterrows():
            text = " | ".join(normalize_text(value) for value in row.tolist() if normalize_text(value))
            if text:
                parts.append(text)
    return "\n".join(parts)


def read_csv_text(path: Path) -> str:
    import pandas as pd

    frame = pd.read_csv(path, dtype=str, encoding="utf-8-sig").fillna("")
    parts = []
    for _, row in frame.head(5000).iterrows():
        text = " | ".join(normalize_text(value) for value in row.tolist() if normalize_text(value))
        if text:
            parts.append(text)
    return "\n".join(parts)


def read_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    text = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(text, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return normalize_text(soup.get_text(" "))


def read_json_text(path: Path) -> str:
    payload = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    return json.dumps(payload, ensure_ascii=False)


def read_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_text(path: Path) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return read_pdf(path), warnings
        if suffix == ".docx":
            return read_docx(path), warnings
        if suffix in {".xlsx", ".xls"}:
            return read_spreadsheet(path), warnings
        if suffix == ".csv":
            return read_csv_text(path), warnings
        if suffix in {".html", ".htm"}:
            return read_html(path), warnings
        if suffix == ".json":
            return read_json_text(path), warnings
        if suffix in {".txt", ".md"}:
            return read_plain_text(path), warnings
    except Exception as exc:
        warnings.append(f"{path.name}: gagal ekstrak {suffix}: {exc}")
        return "", warnings
    warnings.append(f"{path.name}: format {suffix} belum didukung")
    return "", warnings


def term_count(text: str, term: str) -> int:
    if not term:
        return 0
    lower = text.lower()
    if re.fullmatch(r"[a-z0-9]{3,6}(\.jk)?", term.lower()):
        base = term.lower().replace(".jk", "")
        return len(re.findall(rf"(?<![a-z0-9]){re.escape(base)}(?:\.jk)?(?![a-z0-9])", lower))
    return lower.count(term.lower())


def snippets_for_terms(text: str, terms: Sequence[str], limit: int = 3, radius: int = 180) -> List[str]:
    lower = text.lower()
    snippets: List[str] = []
    for term in terms:
        if not term:
            continue
        position = lower.find(term.lower())
        if position < 0:
            continue
        start = max(0, position - radius)
        end = min(len(text), position + len(term) + radius)
        snippet = normalize_text(text[start:end])
        if snippet and snippet not in snippets:
            snippets.append(snippet)
        if len(snippets) >= limit:
            break
    return snippets


def analyze_document(path: Path, profiles: Sequence[Dict[str, str]], min_mentions: int) -> Tuple[List[DocumentHit], List[str]]:
    text, warnings = extract_text(path)
    if not text:
        return [], warnings
    text = normalize_text(text)
    score, label, impact_tags, risk_flags, positives, negatives, base_confidence = score_article(path.name, text[:5000])
    if not impact_tags:
        impact_tags = tags_for_text(text, IMPACT_TAG_TERMS)
    if not risk_flags:
        risk_flags = tags_for_text(text, RISK_FLAG_TERMS)

    hits: List[DocumentHit] = []
    for profile in profiles:
        terms = profile_terms(profile)
        counts = [(term, term_count(text, term)) for term in terms]
        counts = [(term, count) for term, count in counts if count > 0]
        mention_count = sum(count for _, count in counts)
        if mention_count < min_mentions:
            continue
        matched_terms = [term for term, _ in sorted(counts, key=lambda item: item[1], reverse=True)]
        snippets = snippets_for_terms(text, matched_terms)
        confidence = min(1.0, base_confidence + min(0.35, mention_count * 0.03) + (0.08 if snippets else 0.0))
        hits.append(
            DocumentHit(
                ticker=normalize_text(profile.get("ticker")),
                company=normalize_text(profile.get("company")),
                source_file=str(path),
                file_type=path.suffix.lower().lstrip("."),
                mention_count=mention_count,
                matched_terms=";".join(matched_terms[:10]),
                sentiment_score=round(float(score), 3),
                sentiment_label=label,
                impact_tags=";".join(impact_tags[:8]),
                risk_flags=";".join(risk_flags[:8]),
                confidence=round(confidence, 2),
                snippets=" || ".join(snippets),
                processed_at=utc_now_text(),
            )
        )
    return sorted(hits, key=lambda item: (item.mention_count, item.confidence), reverse=True), warnings


def collect_input_files(paths: Sequence[Path], input_dir: Optional[Path], max_files: int) -> List[Path]:
    files: List[Path] = []
    for path in paths:
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES:
            files.append(path)
        elif path.is_dir():
            files.extend(sorted(p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES))
    if input_dir and input_dir.exists():
        files.extend(sorted(p for p in input_dir.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES))
    unique: List[Path] = []
    seen = set()
    for file in files:
        key = str(file.resolve()).lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(file.resolve())
    return unique[:max_files]


def write_csv(path: Path, rows: List[Dict[str, Any]]) -> None:
    fieldnames = list(DocumentHit.__dataclass_fields__.keys())
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def render_html(rows: List[Dict[str, Any]], warnings: List[str], generated_at: str) -> str:
    table_rows = []
    for row in rows[:300]:
        table_rows.append(
            "<tr>"
            f"<td>{html.escape(str(row.get('ticker', '')))}</td>"
            f"<td>{html.escape(str(row.get('company', '')))}</td>"
            f"<td>{html.escape(Path(str(row.get('source_file', ''))).name)}</td>"
            f"<td>{html.escape(str(row.get('mention_count', '')))}</td>"
            f"<td>{html.escape(str(row.get('sentiment_label', '')))} ({html.escape(str(row.get('sentiment_score', '')) )})</td>"
            f"<td>{html.escape(str(row.get('impact_tags', '')))}</td>"
            f"<td>{html.escape(str(row.get('risk_flags', '')))}</td>"
            f"<td>{html.escape(str(row.get('snippets', ''))[:500])}</td>"
            "</tr>"
        )
    warning_html = "".join(f"<li>{html.escape(item)}</li>" for item in warnings) or "<li>Tidak ada warning utama.</li>"
    return f"""<!doctype html>
<html lang="id">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Document Signal Analysis</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 28px; color: #172033; background: #f5f7fb; }}
    main {{ max-width: 1180px; margin: 0 auto; }}
    table {{ width: 100%; border-collapse: collapse; background: #fff; }}
    th, td {{ padding: 9px; border: 1px solid #dbe4ee; text-align: left; font-size: 13px; vertical-align: top; }}
    th {{ background: #eaf1f8; }}
    .note {{ color: #5f6f82; }}
  </style>
</head>
<body>
<main>
  <h1>Document Signal Analysis</h1>
  <p class="note">Generated at {html.escape(generated_at)}. Output ini adalah triage riset, bukan pengganti membaca dokumen asli.</p>
  <table>
    <thead><tr><th>Ticker</th><th>Company</th><th>Source</th><th>Mentions</th><th>Sentiment</th><th>Tags</th><th>Risks</th><th>Snippets</th></tr></thead>
    <tbody>{''.join(table_rows)}</tbody>
  </table>
  <h2>Warnings</h2>
  <ul>{warning_html}</ul>
</main>
</body>
</html>
"""


def write_outputs(rows: List[DocumentHit], warnings: List[str], output_dir: Path, prefix: str) -> Dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    generated_at = utc_now_text()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    row_dicts = [asdict(row) for row in rows]
    payload = {
        "generated_at": generated_at,
        "rows": row_dicts,
        "warnings": warnings,
        "source_posture": "User uploaded documents parsed locally; verify against original files.",
    }
    paths = {
        "csv": output_dir / f"{prefix}_{timestamp}.csv",
        "json": output_dir / f"{prefix}_{timestamp}.json",
        "html": output_dir / f"{prefix}_{timestamp}.html",
        "latest_csv": output_dir / f"{prefix}.csv",
        "latest_json": output_dir / f"{prefix}.json",
        "latest_html": output_dir / f"{prefix}.html",
    }
    write_csv(paths["csv"], row_dicts)
    write_csv(paths["latest_csv"], row_dicts)
    paths["json"].write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    paths["latest_json"].write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    report = render_html(row_dicts, warnings, generated_at)
    paths["html"].write_text(report, encoding="utf-8")
    paths["latest_html"].write_text(report, encoding="utf-8")
    return paths


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Olah PDF/Word/Excel/CSV/TXT/HTML/JSON menjadi sinyal saham.")
    parser.add_argument("--input", type=Path, action="append", default=[], help="File atau folder input. Bisa dipakai berkali-kali.")
    parser.add_argument("--input-dir", type=Path, help="Folder upload yang akan diproses.")
    parser.add_argument("--watchlist", type=Path, default=DEFAULT_WATCHLIST, help="CSV watchlist IDX.")
    parser.add_argument("--output-dir", type=Path, default=ROOT / "outputs" / "document_analysis", help="Folder output.")
    parser.add_argument("--prefix", default="document_analysis", help="Prefix output.")
    parser.add_argument("--min-mentions", type=int, default=1, help="Minimal jumlah mention agar ticker masuk output.")
    parser.add_argument("--max-files", type=int, default=80, help="Batas jumlah file per run.")
    return parser


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()
    if not args.watchlist.exists():
        parser.error(f"Watchlist tidak ditemukan: {args.watchlist}")
    files = collect_input_files(args.input, args.input_dir, args.max_files)
    if not files:
        parser.error("Tidak ada file input yang didukung.")
    profiles = read_watchlist(args.watchlist)
    all_hits: List[DocumentHit] = []
    warnings: List[str] = []
    for file_path in files:
        print(f"[document] processing {file_path}")
        hits, file_warnings = analyze_document(file_path, profiles, max(1, args.min_mentions))
        all_hits.extend(hits)
        warnings.extend(file_warnings)
    paths = write_outputs(all_hits, warnings, args.output_dir, args.prefix)
    print("Outputs created:")
    for kind, path in paths.items():
        print(f"- {kind}: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
